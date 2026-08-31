import re
import pdfplumber
from docx import Document

EMAIL_PATTERN = re.compile(r"[a-z0-9_.+-]+@[a-z0-9-]+\.[a-z0-9-.]+")
PHONE_PATTERN = re.compile(r"(\+?\d[\d\-\s\(\)]{8,}\d)")


def clean_text(text: str) -> str:
    if not text:
        return ""

    text = text.lower()
    text = re.sub(r"\s+", " ", text)
    text = re.sub(r"[^\w\s\.\,\-\+\/\:\@\(\)]", "", text)

    return text.strip()


def extract_text_from_pdf(file) -> str:
    try:
        file.seek(0)
        with pdfplumber.open(file) as pdf:
            text = " ".join(page.extract_text() or "" for page in pdf.pages)
        return clean_text(text)

    except Exception as e:
        raise RuntimeError(f"Error processing PDF: {e}") from e


def extract_text_from_docx(file) -> str:
    try:
        file.seek(0)
        doc = Document(file)
        text = " ".join(para.text for para in doc.paragraphs if para.text)
        return clean_text(text)

    except Exception as e:
        raise RuntimeError(f"Error processing DOCX: {e}") from e


def extract_resume_text(file, filename: str) -> str:
    if not filename:
        raise ValueError("Filename is required")

    filename = filename.lower()

    if filename.endswith(".pdf"):
        return extract_text_from_pdf(file)

    if filename.endswith(".docx"):
        return extract_text_from_docx(file)

    raise ValueError("Unsupported file format. Upload PDF or DOCX only.")


def detect_contact_info(text: str) -> dict:
    return {
        "has_email": bool(EMAIL_PATTERN.search(text)),
        "has_phone": bool(PHONE_PATTERN.search(text)),
    }


def analyze_pdf_structure(file) -> dict:
    file.seek(0)
    with pdfplumber.open(file) as pdf:
        page_count = len(pdf.pages)
        has_images = any(len(page.images) > 0 for page in pdf.pages)
        has_tables = any(len(page.extract_tables()) > 0 for page in pdf.pages)
        word_count = sum(len((page.extract_text() or "").split()) for page in pdf.pages)

    return {
        "page_count": page_count,
        "word_count": word_count,
        "has_images": has_images,
        "has_tables": has_tables,
    }


def analyze_docx_structure(file) -> dict:
    file.seek(0)
    doc = Document(file)

    word_count = sum(len(p.text.split()) for p in doc.paragraphs if p.text)
    has_tables = len(doc.tables) > 0
    has_images = len(doc.inline_shapes) > 0

    return {
        "page_count": None,
        "word_count": word_count,
        "has_images": has_images,
        "has_tables": has_tables,
    }


def analyze_resume_formatting(file, filename: str, text: str, skill_count: int) -> dict:
    filename = filename.lower()

    if filename.endswith(".pdf"):
        result = analyze_pdf_structure(file)
    elif filename.endswith(".docx"):
        result = analyze_docx_structure(file)
    else:
        raise ValueError("Unsupported file format. Upload PDF or DOCX only.")

    result["contact_info"] = detect_contact_info(text)

    density_pct = (
        round((skill_count / result["word_count"]) * 100, 1)
        if result["word_count"]
        else 0.0
    )
    if density_pct >= 6:
        density_label = "Strong"
    elif density_pct >= 3:
        density_label = "Moderate"
    else:
        density_label = "Low"

    result["skill_density"] = {
        "percentage": density_pct,
        "label": density_label,
        "skills_found": skill_count,
        "word_count": result["word_count"],
    }

    warnings = []

    if result["has_images"]:
        warnings.append(
            "Contains embedded images or graphics - many ATS systems can't read text inside images."
        )
    if result["has_tables"]:
        warnings.append(
            "Contains tables - some ATS parsers scramble or drop table content."
        )
    if result["page_count"] and result["page_count"] > 2:
        warnings.append(
            f"{result['page_count']} pages - most ATS/recruiter screens favor 1-2 pages."
        )
    if result["word_count"] < 150:
        warnings.append(
            "Very little extractable text - the resume may be too sparse or mostly image-based."
        )
    if density_pct < 3:
        warnings.append(
            "Low skill keyword density - consider naming more specific tools and technologies."
        )
    if not result["contact_info"]["has_email"]:
        warnings.append("No email address detected.")
    if not result["contact_info"]["has_phone"]:
        warnings.append("No phone number detected.")

    result["warnings"] = warnings
    result["is_ats_friendly"] = len(warnings) == 0

    return result
